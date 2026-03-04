import io
import re
import json
import random
import time
import requests
import os
from typing import List, Dict, Tuple
from PIL import Image
import PyPDF2
import pytesseract
from pdf2image import convert_from_bytes
from groq import Groq

# Optional: DOCX support
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("⚠️ python-docx not installed. DOCX files will not be supported.")

# ============================================================================
# CONFIGURATION
# ============================================================================

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
if not GROQ_API_KEY:
    raise ValueError("GROQ_API_KEY not found in environment variables")

groq_client = Groq(api_key=GROQ_API_KEY)

# ============================================================================
# GROQ API FUNCTIONS
# ============================================================================

def call_groq_api(prompt: str, max_tokens: int = 8000, temperature: float = 0.7) -> str:
    """Call Groq API with error handling."""
    try:
        print(f"🤖 Calling Groq API...")
        
        chat_completion = groq_client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "You are an expert educational researcher and content creator specializing in exam and placement preparation. You create challenging, diverse questions and flashcards that test deep understanding, application, analysis, and problem-solving skills. Avoid repetition, ensure variety in question types, and focus on high-value concepts for competitive exams and job placements."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ],
            model="llama-3.3-70b-versatile",
            temperature=temperature,
            max_tokens=max_tokens,
            top_p=1,
            stream=False
        )
        
        response = chat_completion.choices[0].message.content
        print(f"✅ Groq API response received ({len(response)} chars)")
        return response
        
    except Exception as e:
        print(f"❌ Groq API error: {e}")
        return None


# ============================================================================
# DIFFICULTY-SPECIFIC PROMPT GENERATION
# ============================================================================

def get_difficulty_prompts(difficulty: str = "medium") -> Dict[str, str]:
    """
    Generate difficulty-specific instructions and styles.
    Returns a dictionary with research and quiz prompts for the given difficulty level.
    """
    difficulty = difficulty.lower().strip()
    
    if difficulty == "easy":
        return {
            "research_instructions": """Focus on foundational concepts and basic understanding:
1. Clear, simple definitions with everyday examples
2. Basic principles and fundamental concepts
3. Common use cases and simple applications
4. Avoid complex edge cases and advanced optimizations
5. Include analogies and simple explanations
6. Focus on memorization of core facts""",
            
            "quiz_style": """Generate EXACTLY {num_questions} simple, foundational MCQs suitable for beginners.
Each MCQ must:
- Test basic understanding and recall of core concepts
- Use simple, everyday language
- Include straightforward distractors (common mistakes)
- Have clear, unambiguous correct answers
- Focus on definitions, basic processes, and simple applications
- Avoid complex scenarios, edge cases, or advanced topics
Ensure questions are easy for someone learning the topic for the first time.""",
            
            "temperature": 0.5
        }
    
    elif difficulty == "hard":
        return {
            "research_instructions": """Focus on advanced, nuanced, and complex aspects:
1. Advanced definitions and edge cases
2. Complex algorithms, frameworks with intricate interactions
3. Performance optimizations and critical pitfalls
4. Real-world complex scenarios and limitations
5. Comparisons with related advanced concepts
6. Interview questions testing deep mastery
7. Recent developments and cutting-edge practices""",
            
            "quiz_style": """Generate EXACTLY {num_questions} challenging, advanced MCQs for competitive exams and job placements.
Each MCQ must:
- Test advanced understanding: application, analysis, synthesis, problem-solving
- Use complex, real-world scenarios or edge cases
- Include tricky distractors based on common misconceptions
- Challenge assumptions and require critical thinking
- Combine multiple concepts or test nuanced differences
- Focus on optimization, trade-offs, and best practices
- Avoid obvious or straightforward answers
Ensure questions test mastery and deep expert-level knowledge.""",
            
            "temperature": 0.8
        }
    
    else:  # medium (default)
        return {
            "research_instructions": """Focus on balanced, practical knowledge:
1. Comprehensive definitions with contexts
2. Key principles, methodologies with examples
3. Important practical applications and use cases
4. Some advanced concepts and common challenges
5. Comparisons with related concepts
6. Real-world considerations and best practices
7. Interview-relevant topics""",
            
            "quiz_style": """Generate EXACTLY {num_questions} moderate-level MCQs suitable for intermediate learners.
Each MCQ must:
- Test understanding and application of core concepts
- Include realistic scenarios or practical situations
- Have distractors based on partial understanding
- Require reasoning beyond simple memorization
- Cover both foundational and intermediate aspects
- Balance clarity with some complexity
- Test both knowledge and application skills
Ensure questions are appropriate for someone with basic knowledge seeking deeper understanding.""",
            
            "temperature": 0.7
        }


# ============================================================================
# TOPIC EXTRACTION
# ============================================================================

def extract_all_topics(text: str) -> List[str]:
    """Extract ALL topics from the document."""
    print(f"\n{'='*60}")
    print("🎯 EXTRACTING ALL TOPICS FROM DOCUMENT")
    print(f"{'='*60}\n")
    
    print(f"📄 Document content:\n{text[:500]}...\n")
    
    topics = []
    
    # Method 1: Split by common delimiters
    lines = text.replace('–', '\n').replace('-', '\n').replace(',', '\n').split('\n')
    
    for line in lines:
        line = line.strip()
        
        # Remove prefixes like numbers, bullets
        line = re.sub(r'^\d+[\.\)]\s*', '', line)
        line = re.sub(r'^[•\-*]\s*', '', line)
        
        # Skip empty or very short lines
        if len(line) < 3:
            continue
        
        # Skip lines with only special characters
        if not any(c.isalpha() for c in line):
            continue
        
        # Clean up the line
        line = line.strip('.,;:!? ')
        
        # Add if it looks like a topic
        if len(line) > 2 and len(line) < 200:
            topics.append(line)
    
    # Method 2: Extract capitalized phrases
    capitalized = re.findall(r'\b[A-Z][a-zA-Z\s]+(?:\([A-Z]+\))?', text)
    for cap in capitalized:
        cap = cap.strip()
        if 3 < len(cap) < 100 and cap not in topics:
            topics.append(cap)
    
    # Remove duplicates while preserving order
    unique_topics = []
    seen = set()
    for topic in topics:
        topic_lower = topic.lower().strip()
        if topic_lower not in seen and len(topic_lower) > 3:
            seen.add(topic_lower)
            unique_topics.append(topic.strip())
    
    print(f"✅ Extracted {len(unique_topics)} unique topics:")
    for i, topic in enumerate(unique_topics, 1):
        print(f"   {i}. {topic}")
    
    return unique_topics


# ============================================================================
# INDIVIDUAL TOPIC RESEARCH & QUIZ GENERATION
# ============================================================================

def research_and_generate_questions_for_topic(topic: str, num_questions: int = 2, difficulty: str = "medium") -> List[Dict]:
    """
    Research a single topic and generate questions from that research.
    This ensures questions are meaningful and based on actual knowledge.
    Difficulty can be: easy, medium, or hard
    """
    print(f"\n{'='*60}")
    print(f"🔬 PROCESSING TOPIC: {topic} (Difficulty: {difficulty})")
    print(f"{'='*60}\n")
    
    # Get difficulty-specific prompts
    difficulty_info = get_difficulty_prompts(difficulty)
    research_instructions = difficulty_info["research_instructions"]
    quiz_style = difficulty_info["quiz_style"]
    temperature = difficulty_info["temperature"]
    
    # Step 1: Research the topic thoroughly
    print(f"📚 Step 1: Researching '{topic}' at {difficulty} level...")
    
    research_prompt = f"""You are an expert in exam and placement preparation. Provide research on the topic: {topic}.

{research_instructions}

Provide 600-800 words of structured information suitable for generating {difficulty} level MCQs."""

    research_content = call_groq_api(research_prompt, max_tokens=3000, temperature=0.6)
    
    if not research_content:
        print(f"⚠️ Research failed for '{topic}', using fallback")
        research_content = f"{topic} is an important concept in the field. It involves various aspects and applications that are crucial for understanding the subject matter."
    
    print(f"✅ Research complete: {len(research_content)} characters")
    print(f"Preview: {research_content[:200]}...\n")
    
    # Step 2: Generate questions from the research
    print(f"🧠 Step 2: Generating {num_questions} {difficulty} questions from research...")
    
    quiz_prompt = f"""Using the following research on {topic}, {quiz_style.format(num_questions=num_questions)}

RESEARCH:
{research_content}

Each MCQ must:
- Have 4 options: 1 correct, 3 distractors
- Include a detailed explanation referencing the research
- Vary question types: no repetition in style or focus

Return ONLY valid JSON:
{{
  "questions": [
    {{
      "question": "Question text for {topic}",
      "options": ["Correct answer", "Distractor 1", "Distractor 2", "Distractor 3"],
      "correct_answer": 0,
      "explanation": "Detailed explanation with references to research",
      "topic": "{topic}"
    }}
  ]
}}"""

    quiz_response = call_groq_api(quiz_prompt, max_tokens=4000, temperature=temperature)
    
    questions = []
    
    if quiz_response:
        try:
            # Extract JSON
            json_match = re.search(r'\{.*\}', quiz_response, re.DOTALL)
            if json_match:
                quiz_data = json.loads(json_match.group())
                questions = quiz_data.get("questions", [])
                
                # Validate questions
                valid_questions = []
                for q in questions:
                    if all(key in q for key in ["question", "options", "correct_answer"]):
                        # Ensure exactly 4 options
                        if len(q["options"]) >= 4:
                            q["options"] = q["options"][:4]
                        else:
                            while len(q["options"]) < 4:
                                q["options"].append(f"Additional option {len(q['options']) + 1}")
                        
                        # Validate correct_answer
                        if isinstance(q["correct_answer"], int) and 0 <= q["correct_answer"] < 4:
                            q["topic"] = topic
                            if "explanation" not in q:
                                q["explanation"] = f"This is the correct answer about {topic}."
                            valid_questions.append(q)
                
                questions = valid_questions
                print(f"✅ Generated {len(questions)} valid questions for '{topic}'")
                
        except json.JSONDecodeError as e:
            print(f"⚠️ JSON parsing error: {e}")
    
    # Fallback if generation failed
    if not questions:
        print(f"⚠️ Using fallback question generation for '{topic}'")
        questions = [{
            "question": f"What is a key challenge in applying {topic} in real-world scenarios?",
            "options": [
                f"Handling edge cases and optimizations in {topic}",
                f"{topic} has no challenges",
                f"{topic} is always straightforward",
                "Ignore {topic} in practice"
            ],
            "correct_answer": 0,
            "explanation": f"Research shows {topic} involves complex edge cases.",
            "topic": topic
        }]
    
    return questions


def research_and_generate_flashcards_for_topic(topic: str, num_cards: int = 2) -> List[Dict]:
    """
    Research a single topic and generate flashcards from that research.
    """
    print(f"\n{'='*60}")
    print(f"📚 GENERATING FLASHCARDS FOR: {topic}")
    print(f"{'='*60}\n")
    
    # Step 1: Research the topic
    print(f"🔬 Step 1: Researching '{topic}'...")
    
    research_prompt = f"""Provide exam-focused research on {topic} for placement and competitive prep.

Include:
1. Core definitions with nuances
2. Key differences from similar concepts
3. Applications, pros/cons, and optimizations
4. Common exam traps and interview questions
5. Mnemonics or memory aids if applicable
6. Recent trends or advanced topics

Deliver 600-800 words of structured, memorable content for flashcards."""

    research_content = call_groq_api(research_prompt, max_tokens=3000, temperature=0.6)
    
    if not research_content:
        research_content = f"{topic} is an important area of study with various applications and considerations."
    
    print(f"✅ Research complete: {len(research_content)} characters\n")
    
    # Step 2: Generate flashcards
    print(f"🎴 Step 2: Generating {num_cards} flashcards...")
    
    flashcard_prompt = f"""From the research on {topic}, create EXACTLY {num_cards} diverse flashcards for exam/placement prep.

RESEARCH:
{research_content}

Each flashcard must:
- Front: A key concept, term, difference, or scenario question
- Back: Concise yet detailed answer with examples, pros/cons, or steps
- Focus on memorization, quick recall, and application
- Vary types: definitions, comparisons, processes, pitfalls
- Differ from MCQs: no options, emphasize key facts/differences

Return ONLY valid JSON:
{{
  "flashcards": [
    {{
      "front": "Key term or question (e.g., 'What are the key differences between {topic} and related concept?')",
      "back": "Detailed, memorable answer with examples and tips",
      "question": "Key term or question",
      "answer": "Detailed, memorable answer with examples and tips",
      "topic": "{topic}"
    }}
  ]
}}

Ensure no overlap with MCQ styles; prioritize unique, high-yield content."""

    flashcard_response = call_groq_api(flashcard_prompt, max_tokens=4000, temperature=0.8)
    
    flashcards = []
    
    if flashcard_response:
        try:
            json_match = re.search(r'\{.*\}', flashcard_response, re.DOTALL)
            if json_match:
                data = json.loads(json_match.group())
                flashcards = data.get("flashcards", [])
                
                # Validate
                valid_cards = []
                for card in flashcards:
                    if "front" in card and "back" in card:
                        card["question"] = card.get("question", card["front"])
                        card["answer"] = card.get("answer", card["back"])
                        card["topic"] = topic
                        valid_cards.append(card)
                
                flashcards = valid_cards
                print(f"✅ Generated {len(flashcards)} valid flashcards for '{topic}'")
                
        except json.JSONDecodeError as e:
            print(f"⚠️ JSON parsing error: {e}")
    
    # Fallback
    if not flashcards:
        print(f"⚠️ Using fallback flashcard for '{topic}'")
        flashcards = [{
            "front": f"Key differences in {topic}?",
            "back": f"{topic} differs from similar concepts in applications and optimizations.",
            "question": f"Key differences in {topic}?",
            "answer": f"{topic} differs from similar concepts in applications and optimizations.",
            "topic": topic
        }]
    
    return flashcards


# ============================================================================
# MAIN GENERATION FUNCTIONS
# ============================================================================

def generate_mcq_quiz(text: str, num_questions: int = 10, difficulty: str = "medium") -> Dict:
    """
    Generate quiz by:
    1. Extracting all topics from document
    2. For each topic: research it and generate questions based on difficulty
    3. Combine all questions
    
    Difficulty levels: easy, medium, hard
    """
    print(f"\n{'='*70}")
    print(f"🎓 STARTING QUIZ GENERATION ({num_questions} questions at {difficulty} level)")
    print(f"{'='*70}\n")
    
    # Extract all topics
    topics = extract_all_topics(text)
    
    if not topics:
        print("⚠️ No topics found!")
        return {"questions": [{
            "question": "No topics found in document",
            "options": ["Please provide a document with clear topics", "Option 2", "Option 3", "Option 4"],
            "correct_answer": 0,
            "explanation": "Document should contain topic names",
            "topic": "Error"
        }]}
    
    print(f"\n📊 Will generate questions from {len(topics)} topics")
    print(f"Target: {num_questions} total questions at {difficulty} level\n")
    
    # Calculate questions per topic
    questions_per_topic = max(1, num_questions // len(topics))
    extra_questions = num_questions % len(topics)
    
    all_questions = []
    
    # Generate questions for each topic
    for i, topic in enumerate(topics):
        # Give some topics extra questions if needed
        topic_question_count = questions_per_topic + (1 if i < extra_questions else 0)
        
        if len(all_questions) >= num_questions:
            break
        
        print(f"\n[{i+1}/{len(topics)}] Processing topic: {topic}")
        print(f"Generating {topic_question_count} {difficulty} question(s)...")
        
        topic_questions = research_and_generate_questions_for_topic(topic, topic_question_count, difficulty=difficulty)
        all_questions.extend(topic_questions)
        
        # Small delay to avoid rate limiting
        if i < len(topics) - 1:
            time.sleep(1)
    
    # Trim to exact number requested
    all_questions = all_questions[:num_questions]
    
    print(f"\n{'='*70}")
    print(f"✅ QUIZ GENERATION COMPLETE!")
    print(f"{'='*70}")
    print(f"Total questions generated: {len(all_questions)}")
    print(f"\n📋 Questions by topic:")
    
    topic_counts = {}
    for q in all_questions:
        topic = q.get("topic", "Unknown")
        topic_counts[topic] = topic_counts.get(topic, 0) + 1
    
    for topic, count in topic_counts.items():
        print(f"   ✓ {topic}: {count} question(s)")
    
    return {"questions": all_questions}


def generate_flashcards(text: str, num_cards: int = 10) -> Dict:
    """
    Generate flashcards by:
    1. Extracting all topics from document
    2. For each topic: research it and generate flashcards
    3. Combine all flashcards
    """
    print(f"\n{'='*70}")
    print(f"📚 STARTING FLASHCARD GENERATION ({num_cards} cards)")
    print(f"{'='*70}\n")
    
    # Extract all topics
    topics = extract_all_topics(text)
    
    if not topics:
        return {"flashcards": [{
            "front": "No topics found",
            "back": "Please provide a document with clear topics",
            "question": "No topics found",
            "answer": "Please provide a document with clear topics",
            "topic": "Error"
        }]}
    
    print(f"\n📊 Will generate flashcards from {len(topics)} topics")
    print(f"Target: {num_cards} total flashcards\n")
    
    # Calculate cards per topic
    cards_per_topic = max(1, num_cards // len(topics))
    extra_cards = num_cards % len(topics)
    
    all_flashcards = []
    
    # Generate flashcards for each topic
    for i, topic in enumerate(topics):
        topic_card_count = cards_per_topic + (1 if i < extra_cards else 0)
        
        if len(all_flashcards) >= num_cards:
            break
        
        print(f"\n[{i+1}/{len(topics)}] Processing topic: {topic}")
        print(f"Generating {topic_card_count} flashcard(s)...")
        
        topic_cards = research_and_generate_flashcards_for_topic(topic, topic_card_count)
        all_flashcards.extend(topic_cards)
        
        # Small delay
        if i < len(topics) - 1:
            time.sleep(1)
    
    # Trim to exact number
    all_flashcards = all_flashcards[:num_cards]
    
    print(f"\n{'='*70}")
    print(f"✅ FLASHCARD GENERATION COMPLETE!")
    print(f"{'='*70}")
    print(f"Total flashcards generated: {len(all_flashcards)}")
    print(f"\n📋 Flashcards by topic:")
    
    topic_counts = {}
    for card in all_flashcards:
        topic = card.get("topic", "Unknown")
        topic_counts[topic] = topic_counts.get(topic, 0) + 1
    
    for topic, count in topic_counts.items():
        print(f"   ✓ {topic}: {count} card(s)")
    
    return {"flashcards": all_flashcards}


# ============================================================================
# TEXT EXTRACTION FUNCTIONS
# ============================================================================

def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract ALL text from PDF."""
    text = ""
    
    try:
        print("📄 Extracting text from PDF...")
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        
        for i, page in enumerate(reader.pages):
            try:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            except Exception as e:
                print(f"  Page {i+1} error: {e}")
        
        if len(text.strip()) < 100:
            print("📸 Attempting OCR...")
            try:
                images = convert_from_bytes(pdf_bytes, dpi=300)
                for img in images:
                    text += pytesseract.image_to_string(img, lang='eng') + "\n"
            except Exception as ocr_error:
                print(f"❌ OCR failed: {ocr_error}")
        
        return text.strip()
        
    except Exception as e:
        print(f"❌ PDF extraction error: {e}")
        return ""


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from DOCX."""
    if not DOCX_SUPPORT:
        raise Exception("DOCX support not available")
    
    try:
        doc = Document(io.BytesIO(docx_bytes))
        text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return text.strip()
        
    except Exception as e:
        print(f"❌ DOCX error: {e}")
        return ""


def extract_text_from_image(image_bytes: bytes) -> str:
    """Extract text from image."""
    try:
        img = Image.open(io.BytesIO(image_bytes))
        if img.mode not in ('RGB', 'L'):
            img = img.convert('RGB')
        text = pytesseract.image_to_string(img, lang='eng')
        return text.strip()
    except Exception as e:
        print(f"❌ OCR error: {e}")
        return ""


def aggressive_ocr_cleanup(text: str) -> str:
    """Clean OCR errors."""
    replacements = {
        r'[|]UMAN': 'HUMAN',
        r'\|UMAN': 'HUMAN',
        r'\s+': ' ',
        r'\n\s*\n\s*\n+': '\n\n'
    }
    
    for pattern, replacement in replacements.items():
        text = re.sub(pattern, replacement, text)
    
    return text.strip()


# ============================================================================
# PUBLIC API
# ============================================================================

def process_document(file_bytes: bytes, filename: str, num_questions: int = 10, num_cards: int = 10):
    """Process document and generate quiz/flashcards."""
    
    # Extract text
    file_type = filename.lower().split('.')[-1]
    
    if file_type == 'pdf':
        text = extract_text_from_pdf(file_bytes)
    elif file_type == 'docx':
        text = extract_text_from_docx(file_bytes)
    elif file_type in ['png', 'jpg', 'jpeg', 'bmp', 'tiff', 'gif']:
        text = extract_text_from_image(file_bytes)
    elif file_type == 'txt':
        text = file_bytes.decode('utf-8', errors='ignore')
    else:
        return {"error": f"Unsupported file type: {file_type}"}
    
    if not text or len(text.strip()) < 10:
        return {
            "error": "Could not extract text from document",
            "details": f"Only {len(text.strip())} characters extracted"
        }
    
    text = aggressive_ocr_cleanup(text)
    
    quiz_data = generate_mcq_quiz(text, num_questions)
    flashcard_data = generate_flashcards(text, num_cards)
    
    return {
        "text": text,
        "full_text_length": len(text),
        "quiz": quiz_data,
        "flashcards": flashcard_data,
        "success": True
    }